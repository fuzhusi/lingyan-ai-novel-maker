"""短篇导出路由：TXT / DOCX / MD / HTML / EPUB。"""
import io
from flask import send_file
from app.models import ShortStory
from app.routes.short_story import short_story_bp


@short_story_bp.route("/<int:story_id>/export/txt")
def export_txt(story_id):
    """Export short story as TXT."""
    story = ShortStory.query.get_or_404(story_id)
    content = story.content or ""
    title = story.title or "无题"

    text = f"{title}\n{'=' * len(title)}\n\n{content}"
    buf = io.BytesIO(text.encode("utf-8"))
    buf.seek(0)
    return send_file(buf, mimetype="text/plain; charset=utf-8",
                     as_attachment=True, download_name=f"{title}.txt")


@short_story_bp.route("/<int:story_id>/export/docx")
def export_docx(story_id):
    """Export short story as DOCX."""
    story = ShortStory.query.get_or_404(story_id)
    content = story.content or ""
    title = story.title or "无题"

    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Crimson Pro'
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.8
        style.paragraph_format.first_line_indent = Cm(0.75)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.first_line_indent = Cm(0)
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(28)

        doc.add_page_break()

        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            doc.add_paragraph(para_text)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True, download_name=f"{title}.docx",
        )
    except Exception as e:
        buf = io.BytesIO(f"导出失败: {e}\n\n{content}".encode("utf-8"))
        buf.seek(0)
        return send_file(buf, mimetype="text/plain; charset=utf-8",
                         as_attachment=True, download_name=f"{title}.txt")


@short_story_bp.route("/<int:story_id>/export/md")
def export_md(story_id):
    """导出短篇为 Markdown 格式"""
    story = ShortStory.query.get_or_404(story_id)
    content = story.content or ""
    title = story.title or "无题"

    lines = [f"# {title}", ""]
    if story.genre:
        lines.append(f"**体裁：** {story.genre}")
    if story.tone:
        lines.append(f"**基调：** {story.tone}")
    if story.genre or story.tone:
        lines.append("")
    lines.append("---")
    lines.append("")
    lines.append(content)
    lines.append("")

    text = "\n".join(lines)
    buf = io.BytesIO(text.encode("utf-8"))
    buf.seek(0)
    return send_file(buf, mimetype="text/markdown; charset=utf-8",
                     as_attachment=True, download_name=f"{title}.md")


@short_story_bp.route("/<int:story_id>/export/html")
def export_html(story_id):
    """导出短篇为 HTML 格式"""
    story = ShortStory.query.get_or_404(story_id)
    content = story.content or ""
    title = story.title or "无题"

    paragraphs = []
    for para in content.split("\n"):
        para = para.strip()
        if para:
            paragraphs.append(f"<p>{para}</p>")

    meta_parts = []
    if story.genre:
        meta_parts.append(f"<span>体裁：{story.genre}</span>")
    if story.tone:
        meta_parts.append(f"<span>基调：{story.tone}</span>")
    meta_html = " | ".join(meta_parts)

    html_content = (
        "<!DOCTYPE html>"
        "<html lang='zh-CN'><head><meta charset='UTF-8'>"
        f"<title>{title}</title>"
        "<style>"
        "body { font-family: 'Georgia', serif; max-width: 700px; margin: 2rem auto; padding: 0 1.5rem; line-height: 2; color: #333; }"
        "h1 { text-align: center; margin-bottom: 2rem; font-size: 2rem; }"
        ".meta { text-align: center; color: #888; font-size: 0.9em; margin-bottom: 2rem; }"
        "p { text-indent: 2em; margin: 0.8em 0; }"
        "</style></head><body>"
        f"<h1>{title}</h1>"
        f"<div class='meta'>{meta_html}</div>"
        + "".join(paragraphs) +
        "</body></html>"
    )

    buf = io.BytesIO(html_content.encode("utf-8"))
    buf.seek(0)
    return send_file(buf, mimetype="text/html; charset=utf-8",
                     as_attachment=True, download_name=f"{title}.html")


@short_story_bp.route("/<int:story_id>/export/epub")
def export_epub(story_id):
    """导出短篇为 EPUB 格式"""
    story = ShortStory.query.get_or_404(story_id)
    content = story.content or ""
    title = story.title or "无题"

    try:
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier(f"lingyan-short-{story_id}")
        book.set_title(title)
        book.set_language("zh")
        book.add_author("灵砚 AI")

        style = b'body { font-family: Georgia, serif; line-height: 1.8; margin: 1em; } h1 { text-align: center; margin-bottom: 2em; } p { text-indent: 2em; margin: 0.5em 0; }'
        css = epub.EpubItem(uid="style", file_name="style/default.css",
                            media_type="text/css", content=style)
        book.add_item(css)

        chapter = epub.EpubHtml(title=title, file_name="story.xhtml", lang="zh")
        paragraphs = []
        for para in content.split("\n"):
            para = para.strip()
            if para:
                paragraphs.append(f"<p>{para}</p>")
        chapter.content = f"<h1>{title}</h1>{''.join(paragraphs)}".encode("utf-8")
        chapter.add_item(css)
        book.add_item(chapter)

        book.toc = [epub.Link("story.xhtml", title, "story")]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ["nav", chapter]

        buf = io.BytesIO()
        epub.write_epub(buf, book)
        buf.seek(0)
        return send_file(buf, mimetype="application/epub+zip",
                         as_attachment=True, download_name=f"{title}.epub")
    except ImportError:
        buf = io.BytesIO("导出 EPUB 需要安装 ebooklib: pip install ebooklib".encode("utf-8"))
        buf.seek(0)
        return send_file(buf, mimetype="text/plain; charset=utf-8",
                         as_attachment=True, download_name=f"{title}.txt")
    except Exception as e:
        buf = io.BytesIO(f"导出失败: {e}\n\n{content}".encode("utf-8"))
        buf.seek(0)
        return send_file(buf, mimetype="text/plain; charset=utf-8",
                         as_attachment=True, download_name=f"{title}.txt")
