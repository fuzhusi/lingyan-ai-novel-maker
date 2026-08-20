import io
from flask import Blueprint, Response, send_file, request
from app.models import db, Novel, Chapter, ChapterVersion

export_bp = Blueprint("export", __name__, url_prefix="/novel/<int:novel_id>/export")


def _get_approved_content(chapter):
    """Get approved version content, or latest version if none approved."""
    approved = (ChapterVersion.query
                .filter_by(chapter_id=chapter.id, approved=True)
                .order_by(ChapterVersion.version_number.desc()).first())
    if approved:
        return approved.content
    latest = (ChapterVersion.query
              .filter_by(chapter_id=chapter.id)
              .order_by(ChapterVersion.version_number.desc()).first())
    return latest.content if latest else ""


def _get_chapters(novel_id):
    """获取章节列表，支持按参数筛选。"""
    chapters_param = request.args.get("chapters", "")
    if chapters_param:
        # 支持格式: "1,3,5-8,10"
        chapter_numbers = set()
        for part in chapters_param.split(","):
            part = part.strip()
            if "-" in part:
                try:
                    start, end = part.split("-", 1)
                    for i in range(int(start), int(end) + 1):
                        chapter_numbers.add(i)
                except ValueError:
                    pass
            else:
                try:
                    chapter_numbers.add(int(part))
                except ValueError:
                    pass
        
        if chapter_numbers:
            return Chapter.query.filter(
                Chapter.novel_id == novel_id,
                Chapter.chapter_number.in_(chapter_numbers)
            ).order_by(Chapter.chapter_number).all()
    
    return Chapter.query.filter_by(novel_id=novel_id).order_by(Chapter.chapter_number).all()


@export_bp.route("/txt")
def export_txt(novel_id):
    novel = Novel.query.get_or_404(novel_id)
    chapters = _get_chapters(novel_id)

    parts = [f"# {novel.title}\n"]
    if novel.synopsis:
        parts.append(f"{novel.synopsis}\n")
    parts.append("=" * 50 + "\n")

    total_chars = 0
    for ch in chapters:
        content = _get_approved_content(ch)
        if content:
            parts.append(f"\n\n{'=' * 50}")
            parts.append(f"第 {ch.chapter_number} 章 {ch.title or ''}")
            parts.append(f"{'=' * 50}\n")
            parts.append(content)
            total_chars += len(content)

    text = "\n".join(parts)
    buf = io.BytesIO(text.encode("utf-8"))
    buf.seek(0)
    filename = f"{novel.title}.txt"
    return send_file(buf, mimetype="text/plain; charset=utf-8",
                     as_attachment=True, download_name=filename)


@export_bp.route("/docx")
def export_docx(novel_id):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return "python-docx 未安装，请运行: pip install python-docx", 500

    novel = Novel.query.get_or_404(novel_id)
    chapters = _get_chapters(novel_id)

    doc = Document()

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(novel.title)
    run.font.size = Pt(28)
    run.bold = True
    title_para.space_after = Pt(20)

    if novel.synopsis:
        synopsis_para = doc.add_paragraph()
        synopsis_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = synopsis_para.add_run(novel.synopsis)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(120, 120, 120)
        synopsis_para.space_after = Pt(40)

    if novel.genre:
        genre_para = doc.add_paragraph()
        genre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = genre_para.add_run(novel.genre)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    for ch in chapters:
        content = _get_approved_content(ch)
        if not content:
            continue

        # Chapter title
        heading = doc.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(f"第 {ch.chapter_number} 章")
        run.font.size = Pt(18)

        if ch.title:
            subheading = doc.add_paragraph()
            subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subheading.add_run(ch.title)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(100, 100, 100)
            subheading.space_after = Pt(20)

        # Content - split by paragraphs
        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.first_line_indent = Pt(24)
                p.paragraph_format.line_spacing = 1.8

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{novel.title}.docx"
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=filename)


@export_bp.route("/md")
def export_markdown(novel_id):
    """导出 Markdown 格式 (P2-2)。"""
    novel = Novel.query.get_or_404(novel_id)
    chapters = _get_chapters(novel_id)

    lines = [f"# {novel.title}", ""]

    if novel.synopsis:
        lines.append(f"> {novel.synopsis}")
        lines.append("")
    if novel.genre:
        lines.append(f"**类型：** {novel.genre}")
        lines.append("")
    if novel.world_intro:
        lines.append("## 世界观")
        lines.append("")
        lines.append(novel.world_intro)
        lines.append("")
    lines.append("---")
    lines.append("")

    for ch in chapters:
        content = _get_approved_content(ch)
        if not content:
            continue
        lines.append(f"## 第 {ch.chapter_number} 章 {ch.title or ''}")
        lines.append("")
        lines.append(content)
        lines.append("")
        lines.append("---")
        lines.append("")

    text = "\n".join(lines)
    buf = io.BytesIO(text.encode("utf-8"))
    buf.seek(0)
    filename = f"{novel.title}.md"
    return send_file(buf,
                     mimetype="text/markdown; charset=utf-8",
                     as_attachment=True, download_name=filename)


@export_bp.route("/html")
def export_html(novel_id):
    """导出 HTML 格式 (P2-2)。"""
    novel = Novel.query.get_or_404(novel_id)
    chapters = _get_chapters(novel_id)

    parts = [
        "<!DOCTYPE html>",
        "<html lang='zh-CN'><head><meta charset='UTF-8'>",
        f"<title>{novel.title}</title>",
        "<style>",
        "body { font-family: 'Georgia', serif; max-width: 800px; margin: 2rem auto; padding: 0 1rem; line-height: 1.8; }",
        "h1 { text-align: center; border-bottom: 2px solid #333; padding-bottom: 1rem; }",
        "h2 { margin-top: 3rem; border-bottom: 1px solid #ccc; padding-bottom: 0.5rem; }",
        ".synopsis { background: #f5f5f5; padding: 1rem; border-left: 4px solid #333; margin: 1.5rem 0; }",
        ".meta { color: #666; font-size: 0.9em; }",
        "</style></head><body>",
        f"<h1>{novel.title}</h1>",
    ]
    if novel.synopsis:
        parts.append(f"<div class='synopsis'>{novel.synopsis}</div>")
    if novel.genre:
        parts.append(f"<p class='meta'>类型：{novel.genre}</p>")

    for ch in chapters:
        content = _get_approved_content(ch)
        if not content:
            continue
        parts.append(f"<h2>第 {ch.chapter_number} 章 {ch.title or ''}</h2>")
        for para in content.split("\n"):
            para = para.strip()
            if para:
                parts.append(f"<p>{para}</p>")

    parts.append("</body></html>")
    text = "\n".join(parts)
    buf = io.BytesIO(text.encode("utf-8"))
    buf.seek(0)
    filename = f"{novel.title}.html"
    return send_file(buf,
                     mimetype="text/html; charset=utf-8",
                     as_attachment=True, download_name=filename)


@export_bp.route("/epub")
def export_epub(novel_id):
    """导出 EPUB 格式 (P2-2)。"""
    try:
        from ebooklib import epub
    except ImportError:
        return "ebooklib 未安装，请运行: pip install ebooklib", 500

    novel = Novel.query.get_or_404(novel_id)
    chapters = _get_chapters(novel_id)

    book = epub.EpubBook()
    book.set_identifier(f"lingyan-{novel.id}")
    book.set_title(novel.title)
    book.set_language("zh")
    book.add_author("灵砚")

    for ch in chapters:
        content = _get_approved_content(ch)
        if not content:
            continue
        # 转 HTML 段落
        html_content = "<h2>" + (ch.title or f"第{ch.chapter_number}章") + "</h2>"
        for para in content.split("\n"):
            para = para.strip()
            if para:
                html_content += f"<p>{para}</p>"

        epub_chapter = epub.EpubHtml(
            title=f"第{ch.chapter_number}章 {ch.title or ''}",
            file_name=f"chapter_{ch.chapter_number:03d}.xhtml",
            lang="zh",
        )
        epub_chapter.content = html_content
        book.add_item(epub_chapter)
        book.spine.append(epub_chapter)

    book.toc = tuple(book.spine)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    buf = io.BytesIO()
    epub.write_epub(buf, book)
    buf.seek(0)
    filename = f"{novel.title}.epub"
    return send_file(buf,
                     mimetype="application/epub+zip",
                     as_attachment=True, download_name=filename)
